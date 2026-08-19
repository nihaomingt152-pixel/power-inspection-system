#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 报告导出服务。
基于 python-docx 生成包含检测图片、统计表格、AI 分析文本和预警结论的专业报告。
"""

import os
import logging
from pathlib import Path
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# 类别中文名
CLASS_NAMES_ZH = {
    0: "鸟巢", 1: "风筝", 2: "气球", 3: "垃圾",
    4: "绝缘体外壳", 5: "破损的绝缘壳",
    6: "闪燃损坏的绝缘器外壳", 7: "良好的绝缘外壳",
}

# 严重等级对应颜色
SEVERITY_COLORS = {
    "一般": RGBColor(0x22, 0x7C, 0x34),   # 绿色
    "严重": RGBColor(0xFF, 0x8C, 0x00),   # 橙色
    "紧急": RGBColor(0xE0, 0x00, 0x00),   # 红色
}


def export_word_report(record, output_dir: Path) -> str:
    """
    根据检测记录生成 Word 报告。

    Args:
        record: DetectionRecord ORM 对象
        output_dir: 输出目录

    Returns:
        生成的 .docx 文件路径
    """
    doc = Document()

    # --- 设置默认字体为中文字体 ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # --- 封面标题 ---
    title = doc.add_heading("电力输电线路智能检测分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # 空行

    # --- 一、基本信息 ---
    doc.add_heading("一、基本信息", level=1)
    info_table = doc.add_table(rows=7, cols=2, style="Light Shading Accent 1")
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    info_data = [
        ("记录 ID", str(record.id)),
        ("数据来源", record.source_type or "未知"),
        ("文件名", record.source_name or "未知"),
        ("检测时间", record.created_at.strftime("%Y-%m-%d %H:%M:%S") if record.created_at else "未知"),
        ("GPS 坐标", f"{record.gps_latitude}, {record.gps_longitude}" if record.gps_latitude else "无 GPS 信息"),
        ("拍摄时间", record.capture_timestamp.strftime("%Y-%m-%d %H:%M:%S") if record.capture_timestamp else "无时间戳"),
        ("视频时间点", f"{record.frame_timestamp_seconds:.1f} 秒" if record.frame_timestamp_seconds else "不适用"),
    ]
    for i, (key, value) in enumerate(info_data):
        info_table.cell(i, 0).text = key
        info_table.cell(i, 1).text = value

    doc.add_paragraph()

    # --- 二、检测结果概览 ---
    doc.add_heading("二、检测结果概览", level=1)

    # 统计摘要
    summary = doc.add_paragraph()
    total = record.total_detections or 0
    defect_count = record.defect_count or 0
    normal_count = record.normal_count or 0
    summary.add_run(f"共检测到 {total} 个目标").bold = True
    summary.add_run(f"，其中缺陷类 {defect_count} 个，正常类 {normal_count} 个。")

    # 耗时信息
    timing = doc.add_paragraph()
    timing.add_run("处理耗时: ")
    if record.yolo_time_ms:
        timing.add_run(f"YOLO 检测 {record.yolo_time_ms:.0f}ms")
    if record.ai_time_ms:
        timing.add_run(f"，AI 分析 {record.ai_time_ms:.0f}ms")
    if record.total_time_ms:
        timing.add_run(f"，总计 {record.total_time_ms:.0f}ms")

    doc.add_paragraph()

    # --- 三、检测详情表格 ---
    doc.add_heading("三、检测目标详情", level=1)

    detections = record.yolo_detections or []
    if detections:
        det_table = doc.add_table(rows=1, cols=5, style="Table Grid")
        det_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        headers = ["序号", "目标类别", "置信度", "边界框坐标", "分类"]
        hdr_cells = det_table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        for i, d in enumerate(detections, 1):
            row = det_table.add_row()
            bbox = d.get("bbox", {})
            cls_id = d.get("class_id", -1)
            cls_name = d.get("class_name", "未知")

            row.cells[0].text = str(i)
            row.cells[1].text = CLASS_NAMES_ZH.get(cls_id, cls_name)
            row.cells[2].text = f"{d.get('confidence', 0):.2%}"
            row.cells[3].text = (
                f"({bbox.get('x1', 0):.0f}, {bbox.get('y1', 0):.0f}) - "
                f"({bbox.get('x2', 0):.0f}, {bbox.get('y2', 0):.0f})"
            )
            row.cells[4].text = "缺陷" if cls_id in {1, 2, 3, 5, 6} else "正常"
    else:
        doc.add_paragraph("未检测到任何目标。")

    doc.add_paragraph()

    # --- 四、AI 智能分析 ---
    doc.add_heading("四、AI 智能分析报告", level=1)

    ai = record.ai_analysis or {}
    if ai:
        sections = [
            ("缺陷描述", ai.get("description", "无")),
            ("严重程度", ai.get("severity", "未知")),
            ("判级理由", ai.get("severity_reason", "无")),
            ("故障成因", ai.get("cause", "未知")),
            ("维修建议", ai.get("suggestion", "无")),
        ]

        for label, content in sections:
            p = doc.add_paragraph()
            p.add_run(f"【{label}】").bold = True
            p.add_run(f"\n{content}")
    else:
        doc.add_paragraph("未执行 AI 分析。")

    doc.add_paragraph()

    # --- 五、预警结论 ---
    doc.add_heading("五、预警结论", level=1)

    alert_level = record.alert_level or "无"
    alert_msg = record.alert_message or ""

    p = doc.add_paragraph()
    triggered = record.alert_triggered
    if triggered:
        color = SEVERITY_COLORS.get(alert_level, RGBColor(0, 0, 0))
        run = p.add_run(f"⚠ 预警已触发")
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = color
        p.add_run(f"\n预警等级: {alert_level}")
        p.add_run(f"\n详情: {alert_msg}")
    else:
        run = p.add_run("✓ 未触发预警")
        run.font.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = SEVERITY_COLORS["一般"]
        p.add_run(f"\n系统未检测到需要预警的异常情况。")

    doc.add_paragraph()

    # --- 六、附：带标注图片 ---
    doc.add_heading("六、检测标注图片", level=1)
    if record.annotated_image_path and Path(record.annotated_image_path).exists():
        try:
            doc.add_picture(
                str(record.annotated_image_path),
                width=Inches(5.5),
            )
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[图片加载失败: {e}]")
    else:
        doc.add_paragraph("[未保存标注图片]")

    # --- 保存文件 ---
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("—— 报告由电力输电线路智能检测分析预警系统自动生成 ——")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"检测报告_{record.id}_{timestamp}.docx"
    filepath = str(output_dir / filename)
    doc.save(filepath)

    logger.info(f"Word 报告已生成: {filepath}")
    return filepath


# 工单操作日志动作 → 中文标签（导出报告时间线用）
ORDER_ACTION_LABELS = {
    "created": "📝 创建工单",
    "accepted": "✅ 确认接单",
    "submitted": "📤 提交复检",
    "approved": "✔️ 确认闭环",
    "rejected": "❌ 驳回工单",
    "reassigned": "📤 重新派发",
    "deleted": "🗑️ 删除工单",
}


def export_work_order_report(order, db, output_dir: Path) -> str:
    """
    根据闭环工单生成 Word 报告（工单闭环报告）。

    Args:
        order: WorkOrder ORM 对象（须为 closed 状态）
        db: 数据库会话（用于查询操作日志与关联检测记录的 AI 分析）
        output_dir: 输出目录

    Returns:
        生成的 .docx 文件路径
    """
    from database.models import DetectionRecord, OrderLog

    doc = Document()

    # --- 设置默认字体为中文字体（与检测报告保持一致）---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # --- 封面/标题区 ---
    title = doc.add_heading("电力输电线路智能巡检平台", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("工单闭环报告")
    run.font.size = Pt(18)
    run.font.bold = True

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub2.add_run(
        f"工单编号: #{order.id}    |    状态: 已闭环    |    "
        f"生成日期: {datetime.now().strftime('%Y-%m-%d')}"
    )
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # 空行

    # --- 一、工单基本信息 ---
    doc.add_heading("一、工单基本信息", level=1)

    creator_name = order.creator.full_name if order.creator else f"用户#{order.created_by}"
    assignee_name = order.assignee.full_name if order.assignee else f"用户#{order.assigned_to}"

    info_data = [
        ("工单 ID", str(order.id)),
        ("标题", order.title or "无"),
        ("描述", order.description or "无"),
        ("严重等级", order.severity),
        ("状态", "已闭环"),
        ("创建人", creator_name),
        ("指派给", assignee_name),
        ("创建时间", order.created_at.strftime("%Y-%m-%d %H:%M:%S") if order.created_at else "未知"),
        ("闭环时间", order.updated_at.strftime("%Y-%m-%d %H:%M:%S") if order.updated_at else "未知"),
        ("闭环备注", order.close_remark or "无"),
        ("GPS 坐标", f"{order.gps_lat}, {order.gps_lng}"
         if order.gps_lat is not None else "无"),
    ]

    info_table = doc.add_table(rows=len(info_data), cols=2, style="Light Shading Accent 1")
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (key, value) in enumerate(info_data):
        info_table.cell(i, 0).text = key
        info_table.cell(i, 1).text = value

    doc.add_paragraph()

    # --- 二、修复前后对比图片（左右并排）---
    doc.add_heading("二、修复前后对比", level=1)

    def _fill_img_cell(cell, img_path, caption):
        """向表格单元格写入标题 + 图片（图片加载失败时降级为文字提示）。"""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(caption).bold = True
        if img_path and Path(img_path).exists():
            try:
                # run.add_picture 插入后可控制宽度，避免超宽撑破表格
                pic_p = cell.add_paragraph()
                pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic_p.add_run().add_picture(str(img_path), width=Inches(5.5))
            except Exception as e:
                cell.add_paragraph(f"[图片加载失败: {e}]")
        else:
            cell.add_paragraph("[未保存图片]")

    img_table = doc.add_table(rows=1, cols=2, style="Table Grid")
    img_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 优先使用标注图（带 YOLO 框），缺失时回退到原图
    orig_path = order.annotated_image_path or order.original_image_path
    _fill_img_cell(img_table.cell(0, 0), orig_path, "🔍 原始缺陷图（YOLO 标注）")
    _fill_img_cell(img_table.cell(0, 1), order.repair_image_path, "🔧 修复照片（Worker 提交）")

    doc.add_paragraph()

    # --- 三、AI 智能分析报告（来自关联检测记录）---
    doc.add_heading("三、AI 智能分析报告", level=1)

    ai = {}
    if order.detection_record_id:
        rec = db.query(DetectionRecord).filter(
            DetectionRecord.id == order.detection_record_id
        ).first()
        if rec and rec.ai_analysis:
            ai = rec.ai_analysis

    if ai:
        sections = [
            ("严重程度", ai.get("severity", "未知")),
            ("缺陷描述", ai.get("description", "无")),
            ("故障成因推断", ai.get("cause", "未知")),
            ("维修建议", ai.get("suggestion", "无")),
        ]
        for label, content in sections:
            p = doc.add_paragraph()
            p.add_run(f"【{label}】").bold = True
            p.add_run(f"\n{content}")
    else:
        doc.add_paragraph("该工单未关联 AI 分析数据。")

    doc.add_paragraph()

    # --- 四、操作日志时间线 ---
    doc.add_heading("四、操作日志时间线", level=1)

    logs = db.query(OrderLog).filter(
        OrderLog.order_id == order.id
    ).order_by(OrderLog.created_at.asc()).all()

    if logs:
        for lg in logs:
            label = ORDER_ACTION_LABELS.get(lg.action, lg.action)
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{label}").bold = True
            time_str = lg.created_at.strftime("%Y-%m-%d %H:%M:%S") if lg.created_at else ""
            p.add_run(f"　{time_str} — {lg.operator_name}")
            if lg.content:
                p.add_run(f"\n　　{lg.content}")
    else:
        doc.add_paragraph("暂无操作日志。")

    doc.add_paragraph()

    # --- 保存文件 ---
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("—— 报告由电力输电线路智能检测分析预警系统自动生成 ——")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    os.makedirs(output_dir, exist_ok=True)
    # 标题中的非法文件名字符（\/:*?"<>|）需清理，避免保存失败
    safe_title = "".join(c for c in (order.title or "工单") if c not in '\\/:*?"<>|') or "工单"
    date_str = datetime.now().strftime("%Y-%m-%d")
    filename = f"工单报告_#{order.id}_{safe_title}_{date_str}.docx"
    filepath = str(output_dir / filename)
    doc.save(filepath)

    logger.info(f"工单闭环报告已生成: {filepath}")
    return filepath
